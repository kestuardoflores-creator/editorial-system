"""
assembler.py — Sistema de Gestión Editorial
Convierte archivos Markdown en un documento Word final con estilos normativos.

Uso:
    python assembler/assembler.py --normativa apa7
    python assembler/assembler.py --normativa ieee --output MiTesis_Final.docx
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from copy import deepcopy

# ── Imports (python-docx) ──────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import openpyxl
    import bibtexparser
except ImportError:
    print("Instalando dependencias...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install",
                           "python-docx", "openpyxl", "bibtexparser", "Pillow"])
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import openpyxl
    import bibtexparser

# ── Paths ──────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent.parent
MARKDOWNS_DIR = ROOT / "markdowns"
WORD_DIR      = ROOT / "word"
ASSETS_DIR    = ROOT / "assets"
CONFIG_DIR    = ROOT / "config"

WORD_DIR.mkdir(exist_ok=True)

# ── Alignment map ──────────────────────────────────────────────────────────────
ALIGN_MAP = {
    "LEFT":    WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER":  WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT":   WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# ── Heading level map: Markdown prefix → style ID ─────────────────────────────
HEADING_MAP = {
    1: "TIT_CAP",
    2: "H1_APA",
    3: "H2_APA",
    4: "H3_APA",
    5: "H4_APA",
    6: "H5_APA",
}

# ── Callout regex: > [!LABEL attr="val"] ──────────────────────────────────────
CALLOUT_RE  = re.compile(r'^\s*>\s*\[!(\w+)(.*?)\]\s*$')
ATTR_RE     = re.compile(r'(\w+)="([^"]*)"')
CAPTION_RE  = re.compile(r'^\s*>\s+(.+)$')
CITE_RE     = re.compile(r'\[@([^\]]+)\]')

# ─────────────────────────────────────────────────────────────────────────────
# STYLE ENGINE
# ─────────────────────────────────────────────────────────────────────────────

def apply_style(paragraph, style_cfg):
    """Apply a style config dict to a python-docx paragraph."""
    pf = paragraph.paragraph_format

    # Alignment
    align_key = str(style_cfg.get("Alineacion", "LEFT")).upper()
    paragraph.alignment = ALIGN_MAP.get(align_key, WD_ALIGN_PARAGRAPH.LEFT)

    # Spacing
    space_before = style_cfg.get("Espaciado_Antes", 0)
    space_after  = style_cfg.get("Espaciado_Despues", 0)
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)

    # First line indent (positive = indent, negative = hanging)
    sangria = float(style_cfg.get("Sangria_1era", 0))
    if sangria > 0:
        pf.first_line_indent = Pt(sangria)
    elif sangria < 0:
        pf.first_line_indent = Pt(sangria)   # hanging indent (negative)
        pf.left_indent       = Pt(abs(sangria))

    # Line spacing
    interlineado = float(style_cfg.get("Interlineado", 2.0))
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = interlineado

    return paragraph


def apply_run_style(run, style_cfg):
    """Apply font properties to a run."""
    fuente  = style_cfg.get("Fuente", "Times New Roman")
    tamano  = float(style_cfg.get("Tamano", 12))
    negrita = style_cfg.get("Negrita", False)
    italica = style_cfg.get("Italica", False)
    color   = style_cfg.get("Color_Texto", "#000000").lstrip("#")

    run.font.name = fuente
    run.font.size = Pt(tamano)
    run.font.bold   = bool(negrita) if not isinstance(negrita, bool) else negrita
    run.font.italic = bool(italica) if not isinstance(italica, bool) else italica
    if len(color) == 6:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def add_styled_paragraph(doc, text, style_cfg, inline_citations=None, citation_map=None):
    """Add a paragraph with full style applied. Handles inline citations."""
    p = doc.add_paragraph()
    apply_style(p, style_cfg)

    # Process inline citations inside text
    if CITE_RE.search(text) and citation_map:
        _add_text_with_citations(p, text, style_cfg, citation_map)
    else:
        run = p.add_run(text)
        apply_run_style(run, style_cfg)

    return p


def _add_text_with_citations(paragraph, text, style_cfg, citation_map):
    """Split text by [@key] markers and add styled runs with formatted citations."""
    parts = CITE_RE.split(text)
    # CITE_RE.split alternates: text, key, text, key, ...
    i = 0
    while i < len(parts):
        chunk = parts[i]
        if chunk:
            run = paragraph.add_run(chunk)
            apply_run_style(run, style_cfg)
        i += 1
        if i < len(parts):
            key = parts[i]
            citation_text = _format_citation_inline(key, citation_map)
            run = paragraph.add_run(citation_text)
            apply_run_style(run, style_cfg)
            i += 1


def _format_citation_inline(key, citation_map):
    """Format a citation key as APA inline: (Author, Year)."""
    keys = [k.strip().lstrip("@") for k in key.split(";")]
    parts = []
    for k in keys:
        entry = citation_map.get(k)
        if entry:
            author = entry.get("author", "")
            year   = entry.get("year", "n.d.")
            last   = _extract_last_name(author)
            parts.append(f"{last}, {year}")
        else:
            parts.append(k)
    return "(" + "; ".join(parts) + ")"


def _extract_last_name(author_field):
    """Extract first author last name from BibTeX author field."""
    if not author_field:
        return "?"
    first_author = author_field.split(" and ")[0].strip()
    if "," in first_author:
        return first_author.split(",")[0].strip()
    parts = first_author.split()
    return parts[-1] if parts else first_author


# ─────────────────────────────────────────────────────────────────────────────
# PAGE BREAK ENGINE
# ─────────────────────────────────────────────────────────────────────────────

def add_page_break(doc, break_type):
    """
    Add a page break before a chapter.
    NUEVA   → simple page break
    IMPAR   → odd page section break
    PAR     → even page section break
    CONTINUO → nothing
    """
    if break_type == "CONTINUO":
        return

    if break_type == "NUEVA":
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(docx_break_type("page"))
        return

    # IMPAR or PAR: add section break
    section_type = WD_SECTION.ODD_PAGE if break_type == "IMPAR" else WD_SECTION.EVEN_PAGE
    try:
        doc.add_section(section_type)
    except Exception:
        # Fallback to page break
        p = doc.add_paragraph()
        run = p.add_run()
        from docx.oxml import OxmlElement
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)


def docx_break_type(break_name):
    """Return the docx break enum."""
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


# ─────────────────────────────────────────────────────────────────────────────
# MARKDOWN PARSER
# ─────────────────────────────────────────────────────────────────────────────

class ParsedElement:
    """Represents a parsed element from Markdown."""
    def __init__(self, kind, **kwargs):
        self.kind = kind          # "heading" | "paragraph" | "figure" | "table" | "blank"
        self.kwargs = kwargs


def parse_markdown(filepath):
    """
    Parse a Markdown file and return a list of ParsedElement objects.

    Supports:
    - # Heading 1–6
    - Normal paragraphs (→ TEXTO_APA)
    - Callouts: > [!FIG_TIT src="..."] / > [!TABLA_TIT src="..." sheet="..."]
    - Inline citations: [@key] or [@key1; @key2]
    """
    elements = []
    lines = Path(filepath).read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip blank lines
        if not line.strip():
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            elements.append(ParsedElement("heading", level=level, text=text))
            i += 1
            continue

        # Callout start: > [!TAG attr="val"]
        m = CALLOUT_RE.match(line)
        if m:
            tag    = m.group(1).upper()
            attrs  = dict(ATTR_RE.findall(m.group(2)))
            # Next lines starting with > are caption lines
            caption_parts = []
            i += 1
            while i < len(lines) and CAPTION_RE.match(lines[i]):
                caption_parts.append(CAPTION_RE.match(lines[i]).group(1))
                i += 1
            caption = " ".join(caption_parts).strip()
            elements.append(ParsedElement("callout", tag=tag, attrs=attrs, caption=caption))
            continue

        # Normal paragraph (join continuation lines)
        text_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not CALLOUT_RE.match(lines[i]):
            text_lines.append(lines[i])
            i += 1
        text = " ".join(text_lines).strip()
        if text:
            elements.append(ParsedElement("paragraph", text=text))

    return elements


# ─────────────────────────────────────────────────────────────────────────────
# NUMBERING ENGINE
# ─────────────────────────────────────────────────────────────────────────────

# ── Roman numeral helper ──────────────────────────────────────────────────────
def _to_roman(n):
    """Convert integer to uppercase Roman numeral string."""
    val = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
    syms = ["M","CM","D","CD","C","XC","L","XL","X","IX","V","IV","I"]
    result = ""
    for i, v in enumerate(val):
        while n >= v:
            result += syms[i]
            n -= v
    return result


class NumberingEngine:
    """Tracks and builds element numbers per chapter."""

    def __init__(self):
        self.chapter_num  = 0
        self.counters     = {}   # (label_id, chapter_num) → count  (CAPITULO_ELEMENTO)
        self.global_cnt   = {}   # label_id → count                  (CONTINUO)
        self.figures      = []   # [(prefix, caption)]
        self.tables       = []   # [(prefix, caption)]

    def next_chapter(self):
        self.chapter_num += 1

    def _format_n(self, n, formato_numero):
        """Format a number as ARABIC, ROMAN_UPPER, or ROMAN_LOWER."""
        fmt = str(formato_numero).upper() if formato_numero else "ARABIC"
        if fmt == "ROMAN_UPPER":
            return _to_roman(n)
        elif fmt == "ROMAN_LOWER":
            return _to_roman(n).lower()
        return str(n)

    def build_prefix(self, style_cfg):
        """Return numbered prefix string like 'Figura 3' / 'Tabla II' / 'Ec. 3.2'."""
        label_id       = style_cfg.get("ID_Etiqueta", "")
        prefijo        = style_cfg.get("Prefijo_Texto", "")
        separador      = style_cfg.get("Separador_Num", ".")
        formato        = style_cfg.get("Formato_Prefijo", "CONTINUO")
        formato_numero = style_cfg.get("Formato_Numero", "ARABIC")

        if formato == "CAPITULO_ELEMENTO":
            key = (label_id, self.chapter_num)
            self.counters[key] = self.counters.get(key, 0) + 1
            n   = self.counters[key]
            n_str = self._format_n(n, formato_numero)
            return f"{prefijo} {self.chapter_num}{separador}{n_str}".strip()
        else:
            self.global_cnt[label_id] = self.global_cnt.get(label_id, 0) + 1
            n = self.global_cnt[label_id]
            n_str = self._format_n(n, formato_numero)
            sep = separador if separador.strip() else " "
            return f"{prefijo}{sep}{n_str}".strip()


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL READER
# ─────────────────────────────────────────────────────────────────────────────

def read_excel_table(src_path, sheet_name):
    """Read an Excel sheet and return (headers, rows) as lists of strings."""
    full_path = ROOT / src_path
    if not full_path.exists():
        return None, None
    wb  = openpyxl.load_workbook(full_path, data_only=True)
    ws  = wb[sheet_name] if sheet_name in wb.sheetnames else wb.active
    data = [[str(cell.value) if cell.value is not None else "" for cell in row]
            for row in ws.iter_rows()]
    if not data:
        return [], []
    return data[0], data[1:]


def add_excel_table_to_doc(doc, headers, rows, style_map):
    """Insert an Excel table into the Word document."""
    if not headers:
        return

    num_cols = len(headers)
    table    = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h
        for para in hdr_cells[j].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    # Data rows
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, val in enumerate(row_data[:num_cols]):
            row_cells[j].text = val
            for para in row_cells[j].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# BIBTEX READER
# ─────────────────────────────────────────────────────────────────────────────

def load_bib(bib_path):
    """Load .bib file and return a dict keyed by citation key."""
    if not Path(bib_path).exists():
        return {}
    with open(bib_path, encoding="utf-8") as f:
        db = bibtexparser.load(f)
    return {entry["ID"]: entry for entry in db.entries}


def format_reference_apa7(entry):
    """Format a BibTeX entry as APA 7 reference string."""
    etype = entry.get("ENTRYTYPE", "").lower()
    author  = _format_authors_apa(entry.get("author", ""))
    year    = entry.get("year", "s.f.")
    title   = entry.get("title", "")
    journal = entry.get("journal", "")
    volume  = entry.get("volume", "")
    number  = entry.get("number", "")
    pages   = entry.get("pages", "").replace("--", "–")
    publisher = entry.get("publisher", "")
    url     = entry.get("url", "")
    doi     = entry.get("doi", "")

    if etype == "article":
        ref = f"{author} ({year}). {title}. {journal}"
        if volume:
            ref += f", {volume}"
            if number:
                ref += f"({number})"
        if pages:
            ref += f", {pages}"
        ref += "."
    elif etype == "book":
        ref = f"{author} ({year}). {title}. {publisher}."
    elif etype in ("inproceedings", "conference"):
        booktitle = entry.get("booktitle", "")
        ref = f"{author} ({year}). {title}. En {booktitle}"
        if pages:
            ref += f" (pp. {pages})"
        ref += "."
    else:
        ref = f"{author} ({year}). {title}."

    if doi:
        ref += f" https://doi.org/{doi}"
    elif url:
        ref += f" {url}"

    return ref


def _format_authors_apa(author_field):
    """Format BibTeX author field to APA 7: Last, F. M., & Last, F. M."""
    if not author_field:
        return ""
    authors = [a.strip() for a in author_field.split(" and ")]
    formatted = []
    for a in authors:
        if "," in a:
            parts = a.split(",")
            last  = parts[0].strip()
            first = parts[1].strip() if len(parts) > 1 else ""
            initials = " ".join(w[0] + "." for w in first.split() if w)
            formatted.append(f"{last}, {initials}")
        else:
            parts = a.split()
            if parts:
                last = parts[-1]
                initials = " ".join(p[0] + "." for p in parts[:-1] if p)
                formatted.append(f"{last}, {initials}")
    if len(formatted) > 1:
        return ", ".join(formatted[:-1]) + " y " + formatted[-1]
    return formatted[0] if formatted else ""


# ─────────────────────────────────────────────────────────────────────────────
# TOC / INDEX BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def build_toc(doc, toc_entries, style_map):
    """Build TOC as plain text paragraphs with leader dots and page reference."""
    toc_tit_style = style_map.get("TOC_TITULO", {})
    toc_ent_style = style_map.get("TOC_ENTRADA", {})

    p = add_styled_paragraph(doc, "Tabla de Contenido", toc_tit_style)

    for (level, text, _page) in toc_entries:
        indent = "    " * (level - 1)
        entry_text = f"{indent}{text}"
        add_styled_paragraph(doc, entry_text, toc_ent_style)


def build_index(doc, items, title, style_map):
    """Build Lista de Figuras or Lista de Tablas."""
    idx_tit_style = style_map.get("IDX_TITULO", {})
    idx_ent_style = style_map.get("IDX_ENTRADA", {})

    add_styled_paragraph(doc, title, idx_tit_style)

    for (prefix, caption) in items:
        entry_text = f"{prefix}  {caption}"
        add_styled_paragraph(doc, entry_text, idx_ent_style)


# ─────────────────────────────────────────────────────────────────────────────
# PORTADA BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def build_portada(doc, portada_data, style_map):
    """Build portada page from portada.json data."""
    titulo      = portada_data.get("titulo", "Título de la Tesis")
    autor       = portada_data.get("autor", "Autor")
    institucion = portada_data.get("institucion", "Universidad")
    facultad    = portada_data.get("facultad", "Facultad")
    programa    = portada_data.get("programa", "Programa")
    ciudad      = portada_data.get("ciudad", "Ciudad")
    anio        = portada_data.get("anio", "2024")

    p_tit = style_map.get("PORTADA_TITULO", {})
    p_aut = style_map.get("PORTADA_AUTOR", {})
    p_inf = style_map.get("PORTADA_INFO", {})

    # Spacing paragraphs at the top
    for _ in range(4):
        doc.add_paragraph()

    add_styled_paragraph(doc, titulo, p_tit)
    doc.add_paragraph()
    add_styled_paragraph(doc, autor, p_aut)
    doc.add_paragraph()
    add_styled_paragraph(doc, institucion, p_inf)
    add_styled_paragraph(doc, facultad, p_inf)
    add_styled_paragraph(doc, programa, p_inf)
    doc.add_paragraph()
    add_styled_paragraph(doc, f"{ciudad}, {anio}", p_inf)


# ─────────────────────────────────────────────────────────────────────────────
# HEADER / FOOTER
# ─────────────────────────────────────────────────────────────────────────────

def add_page_numbers(section):
    """Add page number to the footer of a section."""
    footer = section.footer
    para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    run = para.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Add page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT MARGINS
# ─────────────────────────────────────────────────────────────────────────────

def set_margins(section, margenes_cm):
    section.top_margin    = Cm(margenes_cm.get("top",    2.54))
    section.bottom_margin = Cm(margenes_cm.get("bottom", 2.54))
    section.left_margin   = Cm(margenes_cm.get("left",   2.54))
    section.right_margin  = Cm(margenes_cm.get("right",  2.54))


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ASSEMBLER
# ─────────────────────────────────────────────────────────────────────────────

def assemble(normativa_name, output_name=None):
    # Load normativa — prefer local Excel (user-editable), fallback to JSON
    xlsx_path = CONFIG_DIR / f"{normativa_name}.xlsx"
    json_path = CONFIG_DIR / f"{normativa_name}.json"

    if xlsx_path.exists():
        from norm_excel import excel_to_dict
        norm = excel_to_dict(xlsx_path)
        print(f"Leyendo normativa desde Excel: {xlsx_path.name}")
    elif json_path.exists():
        with open(json_path, encoding="utf-8") as f:
            norm = json.load(f)
        print(f"Leyendo normativa desde JSON: {json_path.name}")
    else:
        print(f"Normativa no encontrada: {normativa_name}.xlsx / .json")
        sys.exit(1)

    style_map       = {s["ID_Etiqueta"]: s for s in norm["estilos"]}
    inicio_capitulo = norm.get("inicio_capitulo", "NUEVA")
    margenes_cm     = norm.get("margenes_cm", {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54})

    print(f"✅ Normativa: {norm['normativa']} v{norm['version']}")

    # Load BibTeX
    bib_path  = MARKDOWNS_DIR / "referencias.bib"
    bib_map   = load_bib(bib_path)
    print(f"📚 Referencias BibTeX: {len(bib_map)} entradas")

    # Load portada config
    portada_path = MARKDOWNS_DIR / "portada.json"
    portada_data = {}
    if portada_path.exists():
        with open(portada_path, encoding="utf-8") as f:
            portada_data = json.load(f)

    # Discover markdown files in order
    md_files = sorted(
        [f for f in MARKDOWNS_DIR.glob("*.md")],
        key=lambda f: f.name
    )
    print(f"📄 Archivos Markdown encontrados: {len(md_files)}")
    for f in md_files:
        print(f"   {f.name}")

    # Create Word document
    doc = Document()
    section = doc.sections[0]
    set_margins(section, margenes_cm)
    add_page_numbers(section)

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Portada ────────────────────────────────────────────────────────────────
    print("\n📋 Generando portada...")
    build_portada(doc, portada_data, style_map)

    # Page break after portada
    pg_break = doc.add_paragraph()
    pg_run   = pg_break.add_run()
    from docx.enum.text import WD_BREAK
    pg_run.add_break(WD_BREAK.PAGE)

    # Placeholder for TOC (we'll insert it in memory and track entries)
    toc_entries  = []   # [(level, text, page_placeholder)]
    fig_index    = []   # [(prefix, caption)]
    tab_index    = []   # [(prefix, caption)]
    numbering    = NumberingEngine()
    references_used = set()

    # ── Process chapters ───────────────────────────────────────────────────────
    print("\n📝 Procesando capítulos...")

    for md_idx, md_file in enumerate(md_files):
        print(f"   [{md_idx+1}/{len(md_files)}] {md_file.name}")
        elements = parse_markdown(md_file)

        first_element = True
        for elem in elements:

            if elem.kind == "heading":
                level = elem.kwargs["level"]
                text  = elem.kwargs["text"]
                label = HEADING_MAP.get(level, "TEXTO_APA")
                s     = style_map.get(label, style_map.get("TEXTO_APA", {}))

                # Page break before chapter (only for level-1 headings)
                if level == 1:
                    numbering.next_chapter()
                    if md_idx > 0 or not first_element:
                        add_page_break(doc, inicio_capitulo)
                    toc_entries.append((level, text, ""))

                elif level == 2:
                    toc_entries.append((level, text, ""))

                add_styled_paragraph(doc, text, s, citation_map=bib_map)
                first_element = False

            elif elem.kind == "paragraph":
                text = elem.kwargs["text"]
                s    = style_map.get("TEXTO_APA", {})
                add_styled_paragraph(doc, text, s, citation_map=bib_map)

                # Track cited keys
                for key_group in CITE_RE.findall(text):
                    for key in key_group.split(";"):
                        references_used.add(key.strip().lstrip("@"))

            elif elem.kind == "callout":
                tag     = elem.kwargs["tag"]
                attrs   = elem.kwargs["attrs"]
                caption = elem.kwargs["caption"]
                s       = style_map.get(tag, style_map.get("TEXTO_APA", {}))

                # Build numbered prefix
                if s.get("Es_Numerable", False):
                    prefix = numbering.build_prefix(s)
                    full_caption = f"{prefix}\n{caption}"
                else:
                    prefix = ""
                    full_caption = caption

                # Figure with image
                if tag == "FIG_TIT":
                    src = attrs.get("src", "")
                    img_path = ROOT / src
                    if img_path.exists():
                        try:
                            img_p = doc.add_paragraph()
                            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_p.add_run()
                            run.add_picture(str(img_path), width=Inches(5.5))
                        except Exception as e:
                            print(f"     ⚠️  No se pudo insertar imagen: {e}")

                # Table from Excel
                elif tag == "TABLA_TIT":
                    src   = attrs.get("src", "")
                    sheet = attrs.get("sheet", None)
                    headers, rows = read_excel_table(src, sheet)
                    if headers:
                        add_excel_table_to_doc(doc, headers, rows, style_map)
                    else:
                        print(f"     ⚠️  No se encontró la tabla Excel: {src}")

                # Prefix line
                if prefix:
                    title_line = f"{prefix}"
                    p = add_styled_paragraph(doc, title_line, s)
                    # Caption below
                    if caption:
                        cap_style = deepcopy(s)
                        cap_style["Negrita"] = False
                        add_styled_paragraph(doc, caption, cap_style)
                else:
                    add_styled_paragraph(doc, caption or "", s)

                # Register in indexes
                if tag == "FIG_TIT":
                    fig_index.append((prefix, caption))
                elif tag == "TABLA_TIT":
                    tab_index.append((prefix, caption))

                first_element = False

    # ── TOC ────────────────────────────────────────────────────────────────────
    print("\n📑 Generando Tabla de Contenido...")
    # Insert TOC section before chapters — we append at end for now
    # (A proper pre-chapter TOC requires section manipulation; this version appends)
    toc_break = doc.add_paragraph()
    toc_break.add_run().add_break(WD_BREAK.PAGE)
    build_toc(doc, toc_entries, style_map)

    # ── Lista de Figuras ───────────────────────────────────────────────────────
    if fig_index:
        print(f"\n🖼️  Generando Lista de Figuras ({len(fig_index)} figuras)...")
        fig_break = doc.add_paragraph()
        fig_break.add_run().add_break(WD_BREAK.PAGE)
        build_index(doc, fig_index, "Lista de Figuras", style_map)

    # ── Lista de Tablas ────────────────────────────────────────────────────────
    if tab_index:
        print(f"\n📊 Generando Lista de Tablas ({len(tab_index)} tablas)...")
        tab_break = doc.add_paragraph()
        tab_break.add_run().add_break(WD_BREAK.PAGE)
        build_index(doc, tab_index, "Lista de Tablas", style_map)

    # ── Referencias ────────────────────────────────────────────────────────────
    ref_style = style_map.get("REFERENCIA", {})
    ref_tit   = style_map.get("TOC_TITULO", {})

    if bib_map:
        print(f"\n📚 Generando lista de referencias ({len(references_used)} citadas)...")
        ref_break = doc.add_paragraph()
        ref_break.add_run().add_break(WD_BREAK.PAGE)
        add_styled_paragraph(doc, "Referencias", ref_tit)

        # Sort by author last name
        cited_entries = [bib_map[k] for k in sorted(references_used) if k in bib_map]
        uncited = [e for k, e in bib_map.items() if k not in references_used]
        all_refs = sorted(cited_entries + uncited,
                         key=lambda e: _extract_last_name(e.get("author", "")))

        for entry in all_refs:
            ref_text = format_reference_apa7(entry)
            add_styled_paragraph(doc, ref_text, ref_style)

    # ── Save ───────────────────────────────────────────────────────────────────
    out_name = output_name or "Tesis_Final.docx"
    out_path = WORD_DIR / out_name
    doc.save(str(out_path))
    print(f"\n✅ Documento generado: {out_path}")

    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ensamblador Editorial")
    parser.add_argument("--normativa", default="apa7",
                        help="Nombre de la normativa (sin .json). Ej: apa7, ieee")
    parser.add_argument("--output", default=None,
                        help="Nombre del archivo de salida. Ej: MiTesis.docx")
    args = parser.parse_args()

    print("\n" + "="*60)
    print("  Sistema de Gestión Editorial — Ensamblador")
    print("="*60)

    assemble(args.normativa, args.output)
